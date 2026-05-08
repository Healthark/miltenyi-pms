"""
_build_testing_guide.py — Generates docs/Miltenyi-PMS-Testing-Guide.docx.

Run from the repo root or the docs folder:
    python docs/_build_testing_guide.py

The output is a stakeholder-friendly testing guide / application manual
written for non-technical readers. Re-run any time the seed data, page
layout, or workflow changes — the file is regenerated in place.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT_PATH = Path(__file__).resolve().parent / "Miltenyi-PMS-Testing-Guide.docx"


# ── Theming helpers ─────────────────────────────────────────────────

BRAND = RGBColor(0x44, 0x57, 0xCB)   # tailwind brand-ish indigo
MUTED = RGBColor(0x55, 0x55, 0x55)
HINT  = RGBColor(0x95, 0x6F, 0x00)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = BRAND
        run.font.name = "Calibri"
    return h


def para(doc, text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = "Calibri"
    if size is not None:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt(11)
    return p


def bullet(doc, text, *, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def callout(doc, text, *, color=HINT):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run("Tip: ")
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10.5)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = color
    run.font.size = Pt(10.5)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10.5)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            t.rows[r_idx].cells[c_idx].text = val
            for r in t.rows[r_idx].cells[c_idx].paragraphs[0].runs:
                r.font.size = Pt(10.5)
    return t


# ── Document body ───────────────────────────────────────────────────

def build_document() -> Document:
    doc = Document()

    # Cover page-ish header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Miltenyi PMS")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND
    run.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Stakeholder Testing Guide & Application Manual")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED
    run.font.name = "Calibri"

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run(
        "A walkthrough of every screen, every role, and the things we'd "
        "love your eyes on."
    )
    run.italic = True
    run.font.color.rgb = MUTED
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # ── Section 1: Welcome ─────────────────────────────────────────
    heading(doc, "1. Welcome", level=1)
    para(
        doc,
        "Thank you for taking the time to test the Miltenyi Performance "
        "Management System (PMS). This guide is written for non-technical "
        "users — no jargon, no code, just step-by-step walkthroughs for "
        "every screen and every role.",
    )
    para(
        doc,
        "The goal of this round of testing is simple: log in as different "
        "people, walk through real scenarios, and tell us what feels off, "
        "what's missing, what doesn't make sense, or what surprises you. "
        "Your feedback shapes the next iteration.",
    )

    heading(doc, "What this app is for", level=2)
    para(
        doc,
        "Miltenyi PMS captures performance management for staff who work on "
        "Miltenyi projects. It covers three streams that work together:",
    )
    bullet(doc, "Annual Goals — yearly objectives reviewed twice a year (H1, H2).")
    bullet(doc, "Annual Reviews — a yearly performance appraisal with self → mentor → management stages.")
    bullet(doc, "Project Reviews — per-project, per-cycle evaluations written by the project's PM.")

    heading(doc, "Who's involved", level=2)
    para(
        doc,
        "Five distinct roles share the system. Each has its own set of pages "
        "and permissions:",
    )
    table(
        doc,
        ["Role", "Belongs to", "What they do"],
        [
            ["HR · Healthark", "Healthark (MyOrg)", "Full super-admin. Sees and manages everything across the org."],
            ["HR · Miltenyi", "Miltenyi", "Limited admin. Manages users, projects, project reviews, and system settings — but cannot touch Mentor or Healthark HR rows."],
            ["Mentor", "Healthark", "Reviews each mentee's goals, writes annual mentor evaluations, and views (read-only) their mentees' project reviews."],
            ["PM (Project Manager)", "Miltenyi", "Owns one or more projects and writes the per-cycle project review for every team member on those projects."],
            ["Staff", "Miltenyi domain account, mentored by a Healthark mentor", "Submits annual goals, writes the annual self-review, and sees their own project reviews."],
        ],
    )

    # ── Section 2: Getting Started ─────────────────────────────────
    heading(doc, "2. Getting Started", level=1)

    heading(doc, "How to log in", level=2)
    numbered(doc, "Open the demo URL we shared with you in your browser.")
    numbered(doc, "Type the email of the role you want to test in (see the account list below).")
    numbered(doc, "The password for every test account is: password123")
    numbered(doc, "Click Sign In. You'll land on the role's home page.")
    callout(
        doc,
        "Open the app in two browser windows side-by-side (one regular, one "
        "incognito) so you can be logged in as two roles at once and watch "
        "actions flow from one to the other.",
    )

    heading(doc, "Test accounts (password: password123)", level=2)
    para(doc, "HR", bold=True)
    table(
        doc,
        ["Role", "Email", "Name"],
        [
            ["HR · Healthark", "aanya.sharma@healthark.ai", "Aanya Sharma"],
            ["HR · Miltenyi", "werner@miltenyi.com", "Werner Fischer"],
        ],
    )

    para(doc, "Mentors (Healthark)", bold=True)
    table(
        doc,
        ["Email", "Name", "Mentees"],
        [
            ["rahul.verma@healthark.ai", "Rahul Verma", "Aarav Patel, Diya Mehta, Kabir Singh"],
            ["neha.kapoor@healthark.ai", "Neha Kapoor", "Ishaan Joshi, Saanvi Reddy, Ayaan Khan"],
            ["vikram.iyer@healthark.ai", "Vikram Iyer", "Riya Nair, Arjun Gupta, Myra Desai"],
        ],
    )

    para(doc, "PMs (Miltenyi)", bold=True)
    table(
        doc,
        ["Email", "Name", "Function"],
        [
            ["stefan@miltenyi.com", "Stefan Bauer", "R&D"],
            ["helena@miltenyi.com", "Helena Vogel", "Manufacturing"],
            ["markus@miltenyi.com", "Markus Krause", "Commercial"],
            ["brigitte@miltenyi.com", "Brigitte Hoffmann", "R&D"],
        ],
    )

    para(doc, "Staff (Miltenyi domain — all are mentored by a Healthark mentor)", bold=True)
    table(
        doc,
        ["Email", "Name", "Function · Designation", "Mentor"],
        [
            ["aarav.patel@miltenyi.com",   "Aarav Patel",   "R&D · Scientist",          "Rahul Verma"],
            ["diya.mehta@miltenyi.com",    "Diya Mehta",    "R&D · Senior Scientist",   "Rahul Verma"],
            ["kabir.singh@miltenyi.com",   "Kabir Singh",   "R&D · Scientist",          "Rahul Verma"],
            ["ishaan.joshi@miltenyi.com",  "Ishaan Joshi",  "Manufacturing · Senior Scientist", "Neha Kapoor"],
            ["saanvi.reddy@miltenyi.com",  "Saanvi Reddy",  "Manufacturing · Scientist",        "Neha Kapoor"],
            ["ayaan.khan@miltenyi.com",    "Ayaan Khan",    "Manufacturing · Scientist",        "Neha Kapoor"],
            ["riya.nair@miltenyi.com",     "Riya Nair",     "Commercial · Senior Scientist",    "Vikram Iyer"],
            ["arjun.gupta@miltenyi.com",   "Arjun Gupta",   "Commercial · Scientist",           "Vikram Iyer"],
            ["myra.desai@miltenyi.com",    "Myra Desai",    "Commercial · Scientist",           "Vikram Iyer"],
        ],
    )

    callout(
        doc,
        "The instance starts empty — no projects, no goals, no reviews. "
        "You'll create those as part of testing. That's deliberate: we want "
        "to see what feels intuitive when you're starting from scratch.",
    )

    heading(doc, "Two demo helpers turned on for this round", level=2)
    para(
        doc,
        "To let you see every screen end-to-end without waiting for the "
        "real calendar, we've enabled two test-only conveniences in this "
        "demo instance. Both are off in production.",
    )
    bullet(
        doc,
        "H1 / H2 goal-review window: normally H2 self/mentor reviews are "
        "locked until October. For testing, the calendar gate is bypassed "
        "so you can submit BOTH H1 and H2 goal reviews back-to-back in a "
        "single session and see the full mentor-review screen, not just "
        "H1's slice of it.",
    )
    bullet(
        doc,
        "Project quarter cycling (Q1 → Q2 → Q3 → Q4): the active cycle is "
        "always one quarter at a time, so to test all four quarters of "
        "project reviews, HR · Healthark rotates the cycle in System "
        "Settings (see Scenario G). Past-cycle reviews stay visible, so "
        "you can build up a full Q1+Q2+Q3+Q4 history in one session.",
    )

    # ── Section 3: Layout overview ─────────────────────────────────
    heading(doc, "3. Getting around the app", level=1)
    para(
        doc,
        "Once you sign in, every page shares the same layout:",
    )
    bullet(doc, "Top bar: shows the current cycle (e.g. \"Q1 FY26-27\") and the goal half (\"H1 FY26-27\").")
    bullet(doc, "Left sidebar: navigation links for everything you have access to.")
    bullet(doc, "Main area: the page content. Most pages have tabs across the top.")
    bullet(doc, "Top-right user menu: shows your name, role, and a Sign Out option.")
    para(
        doc,
        "Different roles see different sidebar items. If a link is missing, "
        "your role doesn't have access to it — that's intentional. The "
        "table below shows who sees what.",
    )

    table(
        doc,
        ["Sidebar item", "Staff", "Mentor", "PM", "HR · Miltenyi", "HR · Healthark"],
        [
            ["Dashboard",        "✓", "✓", "✓", "✓", "✓"],
            ["Project Reviews",  "✓", "✓", "✓", "✓", "✓"],
            ["Annual Goals",     "✓", "✓", "—", "—", "✓"],
            ["Annual Reviews",   "✓", "✓", "—", "—", "✓"],
            ["My Mentees",       "—", "✓", "—", "—", "✓"],
            ["Admin Panel",      "—", "—", "—", "✓", "✓"],
            ["Profile",          "✓", "✓", "✓", "✓", "✓"],
        ],
    )

    # ── Section 4: Page-by-page tour ───────────────────────────────
    heading(doc, "4. Page-by-page tour", level=1)

    heading(doc, "4.1 Dashboard", level=2)
    para(
        doc,
        "Currently a placeholder for every role. Each role sees a different "
        "headline (e.g. Healthark HR Dashboard, PM Dashboard) with a note "
        "describing what will eventually show up there. Nothing to test "
        "here today — just confirm the right title shows up for the role "
        "you logged in as.",
    )

    heading(doc, "4.2 Profile", level=2)
    para(
        doc,
        "Two cards side by side:",
    )
    bullet(
        doc,
        "Profile Info (left): your HR-controlled details — name, email, "
        "phone, organization, function, designation, mentor (for Staff), "
        "join date. Read-only.",
    )
    bullet(
        doc,
        "Change Password (right): set a new password. Asks for your "
        "current password, then a new one (8+ chars), and confirms via a "
        "popup before saving.",
    )

    heading(doc, "4.3 Project Reviews", level=2)
    para(
        doc,
        "Where per-project, per-cycle reviews are written and read. The "
        "tabs you see here depend on your role.",
    )

    para(doc, "Staff: \"My Reviews\" tab", bold=True)
    bullet(doc, "Lists every project you're on, with one row per cycle.")
    bullet(doc, "Click any row to expand it and see the PM's competency comments + impact statement once they're written.")
    bullet(doc, "Toggle between Cards view and Table view. Filter by cycle, project, PM, or status.")

    para(doc, "PM: \"Primary Evaluation\" tab", bold=True)
    bullet(doc, "Shows every team member on every project where you are the PM.")
    bullet(doc, "Each row is a (employee, project, cycle) review. Click Evaluate to open the modal and write the review.")
    bullet(doc, "The modal shows the role expectations for that employee's function and designation as a reference. Fill 7 competency comments + a performance rating + an impact statement.")
    bullet(doc, "Save Draft to come back later. Submit to finalise — submitted rows lock and the employee can now see them.")

    para(doc, "Mentor: \"Mentees' Reviews\" tab", bold=True)
    bullet(doc, "Read-only view of every project review for any of your mentees.")
    bullet(doc, "Filter by cycle, project, PM, mentee, or status. Useful for catching what's still pending across mentees.")

    para(doc, "HR (either): \"All Reviews\" tab", bold=True)
    bullet(doc, "Read-only view of every project review across the org.")
    bullet(doc, "Same filters as the mentor view.")

    para(doc, "Anyone listed as Secondary on a project: \"Secondary Evaluation\" tab", bold=True)
    bullet(
        doc,
        "Appears only if HR has set you as the Secondary evaluator on at "
        "least one project. Lets you add a short impact statement to a "
        "project review after the PM has submitted theirs.",
    )

    heading(doc, "4.4 Annual Goals", level=2)
    para(
        doc,
        "Where each Staff person captures their yearly objectives, and "
        "mentors approve / track them.",
    )
    para(doc, "Staff: \"My Goals\" tab", bold=True)
    bullet(doc, "Click \"Add Goal\" to create a goal. Title, description, optional success criteria. Save as Draft or Submit for mentor approval.")
    bullet(doc, "Once approved, you'll see the H1 (and later H2) self-review actions on the same goal card.")
    bullet(doc, "Self-review is half-yearly: write a short paragraph on how you progressed against this goal. Mentor reviews it after.")
    bullet(doc, "Demo note: H1 and H2 are BOTH unlocked for this round, so you can fill both halves back-to-back instead of waiting for October.")

    para(doc, "Mentor: \"Team Goals\" tab", bold=True)
    bullet(doc, "Lists every goal owned by your mentees, grouped by mentee.")
    bullet(doc, "Approve / Reject each submitted goal. After H1 / H2 self-reviews land, write your mentor review on the same row.")

    para(doc, "HR · Healthark: \"All Goals\" tab", bold=True)
    bullet(doc, "Read-only view of every goal across the org. Group view: one row per employee, expand to see their goals.")
    bullet(doc, "Filters: Employee (typeable — start typing a name), Year, Status.")

    heading(doc, "4.5 Annual Reviews", level=2)
    para(
        doc,
        "The yearly performance appraisal — three stages: Self → Mentor → "
        "Management.",
    )
    para(doc, "Staff: \"My Reviews\" tab", bold=True)
    bullet(doc, "Shows your past annual reviews and the current cycle's status.")
    bullet(doc, "Click \"Self-Review\" to write this year's overall self-review and self-rating. Save as draft or submit.")
    bullet(doc, "Once your mentor and management have completed their parts, you can read their feedback and the final rating.")

    para(doc, "Mentor: \"Team Review\" tab", bold=True)
    bullet(doc, "Lists every annual review owned by any of your mentees, every year.")
    bullet(doc, "Filter by year, status, or mentee.")
    bullet(doc, "Click Evaluate on a row that's pending mentor input → write your mentor review and rating, save draft or submit.")

    para(doc, "HR · Healthark: \"All Reviews\" tab", bold=True)
    bullet(doc, "Read-only org-wide table.")
    bullet(doc, "Click any row to expand and read the self + mentor narratives side by side.")
    bullet(doc, "Filters: Employee (typeable), Cycle, Status, Function, Designation.")

    heading(doc, "4.6 My Mentees", level=2)
    para(
        doc,
        "Available to Mentor and HR · Healthark.",
    )
    para(doc, "Mentor view", bold=True)
    bullet(doc, "Master grid: one card per mentee with goal stats, the active-cycle annual review status, and project rollups.")
    bullet(doc, "Click a card to drop into the mentee detail page — every goal, every annual review, every project they've been on, all in one place.")

    para(doc, "HR · Healthark view", bold=True)
    bullet(doc, "Grouped \"All Mentor Pairings\" view: one section per Mentor, mentees listed beneath. Useful for sanity-checking who reports to whom.")

    heading(doc, "4.7 Admin Panel", level=2)
    para(
        doc,
        "Available to HR roles only.",
    )

    para(doc, "Users tab", bold=True)
    bullet(doc, "Add / edit / deactivate users.")
    bullet(doc, "HR · Miltenyi cannot create or edit Mentor or HR · Healthark rows — those are protected (you'll see a \"View-only\" tag instead of edit buttons).")
    bullet(doc, "Filter by role, status, function, or designation.")

    para(doc, "Projects tab", bold=True)
    bullet(doc, "Create, edit, soft-delete projects. Set the PM and the optional Secondary evaluator.")
    bullet(doc, "Project Status column: Active or Completed pill. Use the actions on each row to:")
    bullet(doc, "Mark Complete — archives the project and end-dates every active assignment in one go. Future cycles stop generating new pending reviews. Past reviews stay visible as history.", level=1)
    bullet(doc, "Re-open — flips a Completed project back to Active. Doesn't auto-restore assignments — re-add anyone who should be on the team.", level=1)
    bullet(doc, "Inside the project edit modal, the trash icon next to a team member ends that person's assignment (with an Undo toast for 6 seconds).")

    para(doc, "Management Review tab (HR · Healthark only)", bold=True)
    bullet(doc, "Calibration grid for the annual review's third stage.")
    bullet(doc, "Set the management performance rating per row, toggle final-rating visibility, and finalise.")

    para(doc, "System Settings tab", bold=True)
    bullet(doc, "The cycle (Q1, Q2 etc.), submission gates (open / closed), visibility toggles for ratings.")
    bullet(doc, "For this demo, every gate is already open and every visibility toggle is on. Try flipping things to see how the rest of the app reacts.")

    # ── Section 5: Scenarios ───────────────────────────────────────
    heading(doc, "5. Test scenarios", level=1)
    para(
        doc,
        "These are end-to-end stories that touch every role. Run through "
        "them in order if you can — they build on each other.",
    )

    heading(doc, "Scenario A: Set up a project and run a project review", level=2)
    numbered(doc, "Sign in as Aanya Sharma (HR · Healthark) → Admin Panel → Projects → click \"Add Project\".")
    numbered(doc, "Set a name, code (e.g. MIL-PRJ-T01), pick Stefan Bauer as PM, leave Secondary blank for now.")
    numbered(doc, "Add 2–3 R&D Staff (Aarav, Diya, Kabir) as team members, save.")
    numbered(doc, "Sign out and back in as Stefan Bauer (PM) → Project Reviews → Primary Evaluation tab. You should see your team members listed for Q1 FY26-27.")
    numbered(doc, "Click Evaluate on one row. Fill the 7 comment fields, pick a rating, write an impact statement. Save Draft, then re-open and Submit.")
    numbered(doc, "Sign out and back in as Aarav Patel (Staff) → Project Reviews → My Reviews. You should see the row with status \"Reviewed\". Click to expand and read what the PM wrote.")
    callout(
        doc,
        "Things to watch: do the role expectations on the right of the "
        "evaluation modal match what you'd expect for an R&D Scientist? Is "
        "the rating badge colour helpful? Does the staff side feel polished?",
    )

    heading(doc, "Scenario B: Submit a goal and run BOTH H1 + H2 reviews", level=2)
    numbered(doc, "Sign in as Diya Mehta (Staff) → Annual Goals → My Goals → \"Add Goal\".")
    numbered(doc, "Title, description, success criteria. Submit (don't save as draft).")
    numbered(doc, "Sign in as Rahul Verma (Mentor) → Annual Goals → Team Goals. Diya's goal should show up in \"Pending approval\".")
    numbered(doc, "Approve it.")
    numbered(doc, "Back as Diya: the goal now shows the H1 self-review action. Write a short H1 self-review and submit.")
    numbered(doc, "Back as Rahul: write the H1 mentor review for that goal.")
    numbered(doc, "Back as Diya: now the H2 self-review action should be available too (the H1/H2 calendar gate is bypassed for this demo). Submit a short H2 self-review.")
    numbered(doc, "Back as Rahul: write the H2 mentor review. You should now see the full \"completed for the year\" state on the goal.")
    numbered(doc, "Optional — sign in as Aanya (HR · Healthark) → Annual Goals → All Goals. Type \"Diya\" in the Employee filter and confirm her goal shows up with both H1 and H2 reviews recorded.")
    callout(
        doc,
        "Without the H1/H2 bypass you'd have to wait until October to "
        "submit H2 reviews. We turned the bypass on for this demo so you "
        "can validate both halves' UI in one round.",
    )

    heading(doc, "Scenario C: Run an annual review end-to-end", level=2)
    numbered(doc, "Sign in as Kabir Singh (Staff) → Annual Reviews → My Reviews → click \"Self-Review\".")
    numbered(doc, "Write the overall self-review, pick a rating (1–5), Submit.")
    numbered(doc, "Sign in as Rahul Verma (Mentor) → Annual Reviews → Team Review. Kabir's row is now \"Pending Mentor\".")
    numbered(doc, "Click Evaluate on Kabir's row. Write the mentor review, set a rating, Submit.")
    numbered(doc, "Sign in as Aanya (HR · Healthark) → Admin Panel → Management Review. Kabir's row appears in the calibration grid.")
    numbered(doc, "Set a management rating, toggle visibility on, finalise.")
    numbered(doc, "Back as Kabir: the review is now \"Completed\". Click View — you should see the self review, mentor review (text), and the final rating.")

    heading(doc, "Scenario D: Project completion and assignment lifecycle", level=2)
    numbered(doc, "Sign in as Aanya (HR · Healthark) → Admin Panel → Projects.")
    numbered(doc, "Pick the project from Scenario A. Click Mark Complete. Confirm the warning about end-dating assignments.")
    numbered(doc, "Sign in as Stefan (PM) → Project Reviews → Primary Evaluation. The completed project's team should no longer generate pending placeholders.")
    numbered(doc, "Sign in as Aarav (Staff): his My Reviews still shows the historic row from Scenario A.")
    numbered(doc, "As Aanya: re-open the project. Confirm that team members are NOT auto-restored — you'd need to re-add them.")

    heading(doc, "Scenario E: Mid-cycle removal with Undo", level=2)
    numbered(doc, "As Aanya: open any active project's edit modal → click the trash icon next to a team member → confirm the dialog.")
    numbered(doc, "A toast appears with \"Undo\" for 6 seconds. Click Undo and confirm the row flips back to active.")
    numbered(doc, "Repeat without clicking Undo — confirm the row stays end-dated and renders as a greyed-out historical entry below the active members.")

    heading(doc, "Scenario F: HR · Miltenyi limited admin", level=2)
    numbered(doc, "Sign in as Werner Fischer (HR · Miltenyi) → Admin Panel.")
    numbered(doc, "Confirm: only Users, Projects, and System Settings tabs are visible (no Management Review).")
    numbered(doc, "On Users: try to edit Aanya (HR · Healthark) or any Mentor row. The edit/deactivate icons should be missing — \"View-only\" tag instead.")
    numbered(doc, "Try the Add User button — confirm the role dropdown only offers Staff / PM / HR · Miltenyi (no Mentor, no HR · Healthark).")

    heading(doc, "Scenario G: Project reviews across Q1 → Q4 (HR rotates the cycle)", level=2)
    para(
        doc,
        "The system has one active project-review cycle at a time. To see "
        "what a full year of project reviews looks like, HR rotates the "
        "active cycle through Q1 → Q2 → Q3 → Q4 in System Settings. Past-"
        "cycle reviews stay visible, so by the end you'll have a full "
        "history on each project.",
    )
    numbered(doc, "Sign in as Aanya (HR · Healthark) → Admin Panel → System Settings → set Active Cycle to \"Q1 FY26-27\". Save.")
    numbered(doc, "Sign in as Stefan (PM) → Project Reviews → Primary Evaluation. Pick a row and submit a Q1 review (Scenario A pattern).")
    numbered(doc, "Sign back in as Aanya → System Settings → flip Active Cycle to \"Q2 FY26-27\". Save.")
    numbered(doc, "As Stefan: your Primary Evaluation tab now shows fresh placeholders for Q2 alongside the Q1 reviews you submitted (filter by Cycle to see one at a time). Submit a Q2 review.")
    numbered(doc, "Repeat for Q3 and Q4.")
    numbered(doc, "As an affected Staff member (e.g. Aarav Patel): My Reviews now lists one row per cycle, all reviewed. Click any row to read what the PM wrote.")
    callout(
        doc,
        "Things to watch: how does the PM's queue read when several "
        "cycles' reviews stack up? Does the cycle filter feel intuitive? "
        "Are past-quarter reviews easy to find vs. easy to ignore once "
        "they're reviewed?",
    )

    # ── Section 6: What to look for ────────────────────────────────
    heading(doc, "6. What to look for / what to report", level=1)
    para(doc, "When you find anything off, please capture:")
    bullet(doc, "Which role you were signed in as.")
    bullet(doc, "Which page and tab you were on.")
    bullet(doc, "What you did, what you expected, what actually happened.")
    bullet(doc, "A screenshot if possible.")

    para(doc, "Specific things we'd love feedback on:", bold=True)
    bullet(doc, "Does each page's purpose feel obvious within 5 seconds of landing on it?")
    bullet(doc, "Are headers, button labels, and tab names unambiguous?")
    bullet(doc, "Are the right things shown / hidden per role?")
    bullet(doc, "Filters: are the right ones present? Any obvious filter you'd want that's missing?")
    bullet(doc, "Empty states (no data yet): are they helpful or confusing?")
    bullet(doc, "Confirmation dialogs and toasts: clear messaging? Right level of friction?")
    bullet(doc, "Error messages: helpful or technical-sounding?")
    bullet(doc, "Tone: anywhere a label feels too jargon-y or too informal?")
    bullet(doc, "Anywhere the Staff vs Mentor vs PM mental model breaks down?")
    bullet(doc, "Performance: any page that feels slow loading or sluggish to interact with?")

    # ── Section 7: Things deliberately not in scope ────────────────
    heading(doc, "7. What's intentionally NOT in scope yet", level=1)
    para(
        doc,
        "These are areas you may notice that are placeholders or "
        "deliberately out of scope for this round. No need to report them:",
    )
    bullet(doc, "Dashboard widgets — placeholder copy only.")
    bullet(doc, "Email notifications.")
    bullet(doc, "Mobile-specific layouts (please test on desktop / laptop screens).")
    bullet(doc, "Self-service password reset (use the credentials provided).")
    bullet(doc, "Bulk operations (e.g. importing users via CSV).")
    bullet(doc, "Any analytics, exports, or reporting screens.")

    # ── Section 8: Glossary ────────────────────────────────────────
    heading(doc, "8. Glossary", level=1)
    table(
        doc,
        ["Term", "Meaning"],
        [
            ["Cycle", "A performance period — quarterly (Q1–Q4) or half-yearly (H1–H2)."],
            ["FY26-27", "The fiscal year running April 2026 – March 2027."],
            ["Active cycle", "The cycle currently \"open\" for submissions — visible in the top bar."],
            ["Goal cycle", "Always half-yearly (H1 / H2) regardless of project-review cadence."],
            ["PM (Primary Evaluator)", "The Miltenyi project manager who writes the project review for every team member."],
            ["Secondary Evaluator", "Optional senior who adds an impact statement after the PM submits."],
            ["Mentor", "Healthark person who looks after a Staff person across goals + annual reviews."],
            ["Mentee", "A Staff person assigned to a mentor."],
            ["Role expectation", "A reference paragraph per (Function × Designation) that PMs use as a yardstick when writing reviews."],
            ["Performance rating", "A 1–5 score: 1 = beyond expectations, 5 = did not achieve goals."],
            ["Final rating", "The published rating an employee sees, after management calibration."],
        ],
    )

    # ── Footer note ───────────────────────────────────────────────
    para(doc, "")
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run("Thank you for testing.")
    run.italic = True
    run.font.color.rgb = MUTED
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"Wrote: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
